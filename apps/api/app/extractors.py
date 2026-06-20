from __future__ import annotations

import csv
import io
from pathlib import Path

import fitz
from docx import Document
from openpyxl import load_workbook


def extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        with fitz.open(stream=content, filetype="pdf") as document:
            return "\n".join(page.get_text("text") for page in document)
    if suffix == ".docx":
        document = Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)
    if suffix == ".xlsx":
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"[{sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value not in (None, "")]
                if values:
                    parts.append(" | ".join(values))
        workbook.close()
        return "\n".join(parts)
    if suffix == ".csv":
        return content.decode("utf-8-sig", errors="replace")
    if suffix in {".png", ".jpg", ".jpeg"}:
        try:
            import pytesseract
            from PIL import Image
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)
            return text if text.strip() else f"[Image OCR: {filename} - no text extracted]"
        except Exception as e:
            return f"[Image OCR failed for {filename}: {str(e)}]"
    return ""

